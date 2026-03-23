# -*- coding: utf-8 -*-
"""
generar_fotos_word.py
Genera un Word con las fotos del Timemark agrupadas por punto de control.

Dependencias:
    pip install python-docx

Uso:
    python generar_fotos_word.py
"""

import zipfile
import xml.etree.ElementTree as ET
import math
import io
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ─────────────────────────────────────────────
# RUTAS DE ARCHIVOS (relativas al script)
# ─────────────────────────────────────────────
SCRIPT_DIR   = os.path.dirname(os.path.abspath(__file__))
KMZ_TIMEMARK = os.path.join(SCRIPT_DIR, "Timemark_202603222011.kmz")
KMZ_TRAZO    = os.path.join(SCRIPT_DIR, "TRAZO KM 100-110.kmz")
DOCX_GEO     = os.path.join(SCRIPT_DIR, "ZACUALTIPAN MOLANGO - KM 99+960-105+000.docx")
OUTPUT_DOCX  = os.path.join(SCRIPT_DIR, "Fotografias_por_Punto_Control.docx")

# ─────────────────────────────────────────────
# NAMESPACES KML
# ─────────────────────────────────────────────
NS = {
    "kml":   "http://www.opengis.net/kml/2.2",
    "gx":    "http://www.google.com/kml/ext/2.2",
    "atom":  "http://www.w3.org/2005/Atom",
}

# ─────────────────────────────────────────────
# UTILIDADES
# ─────────────────────────────────────────────

def haversine(lat1, lon1, lat2, lon2):
    """Distancia en metros entre dos puntos GPS."""
    R = 6_371_000
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    dphi  = math.radians(lat2 - lat1)
    dlam  = math.radians(lon2 - lon1)
    a = math.sin(dphi/2)**2 + math.cos(phi1)*math.cos(phi2)*math.sin(dlam/2)**2
    return R * 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))


def leer_kml_de_kmz(kmz_path):
    """Devuelve el árbol XML del primer .kml dentro del KMZ."""
    with zipfile.ZipFile(kmz_path, "r") as z:
        kml_names = [n for n in z.namelist() if n.lower().endswith(".kml")]
        if not kml_names:
            raise FileNotFoundError(f"No se encontró .kml en {kmz_path}")
        with z.open(kml_names[0]) as f:
            return ET.parse(f), z.namelist(), kmz_path


# ─────────────────────────────────────────────
# 1. EXTRAER PUNTOS DE CONTROL DEL DOCX
# ─────────────────────────────────────────────

def extraer_puntos_control_docx(docx_path):
    """
    Lee el reporte geológico y extrae todos los cadenamienos con
    'Punto de control' o encabezados de sección tipo 'KM XXX+XXX'.
    Devuelve lista de dicts: {label, km_num}  ordenada por km_num.
    """
    doc = Document(docx_path)
    patron = re.compile(
        r"(?:Punto\s+de\s+control\s+)?[Kk][Mm]\s*(\d{2,3})[+\s](\d{3})",
        re.IGNORECASE
    )
    encontrados = {}
    for para in doc.paragraphs:
        texto = para.text.strip()
        for m in patron.finditer(texto):
            km_int  = int(m.group(1))
            km_dec  = int(m.group(2))
            km_num  = km_int * 1000 + km_dec          # metros desde km 0
            label   = f"KM {km_int}+{km_dec:03d}"
            if km_num not in encontrados:
                encontrados[km_num] = label

    # Ordenar y devolver
    puntos = [{"label": v, "km_num": k} for k, v in sorted(encontrados.items())]
    print(f"  → {len(puntos)} cadenamienos encontrados en el reporte geológico")
    return puntos


# ─────────────────────────────────────────────
# 2. OBTENER COORDENADAS DE PUNTOS DE CONTROL
#    desde el KMZ de TRAZO
# ─────────────────────────────────────────────

def coordenadas_desde_trazo(kmz_trazo, puntos_control):
    """
    Busca en el KMZ de TRAZO los placemarks cuyo nombre coincida con
    el cadenamiento de cada punto de control.
    Si no encuentra coincidencia exacta, interpola linealmente sobre
    la polilínea del eje (LineString más larga).
    """
    tree, _, _ = leer_kml_de_kmz(kmz_trazo)
    root = tree.getroot()

    # ── a) Recopilar todos los placemarks con coordenada puntual ──
    pm_coords = {}   # nombre → (lat, lon)
    for pm in root.iter("{http://www.opengis.net/kml/2.2}Placemark"):
        name_el = pm.find("{http://www.opengis.net/kml/2.2}name")
        point_el = pm.find(".//{http://www.opengis.net/kml/2.2}Point/"
                           "{http://www.opengis.net/kml/2.2}coordinates")
        if name_el is not None and point_el is not None:
            raw = point_el.text.strip().split(",")
            if len(raw) >= 2:
                lon, lat = float(raw[0]), float(raw[1])
                pm_coords[name_el.text.strip()] = (lat, lon)

    # ── b) Recopilar la polilínea del eje (LineString más larga) ──
    eje_coords = []
    max_pts = 0
    for ls in root.iter("{http://www.opengis.net/kml/2.2}LineString"):
        coord_el = ls.find("{http://www.opengis.net/kml/2.2}coordinates")
        if coord_el is not None and coord_el.text:
            pts = []
            for token in coord_el.text.strip().split():
                parts = token.split(",")
                if len(parts) >= 2:
                    pts.append((float(parts[1]), float(parts[0])))  # (lat, lon)
            if len(pts) > max_pts:
                max_pts = len(pts)
                eje_coords = pts

    print(f"  → Placemarks puntuales en TRAZO: {len(pm_coords)}")
    print(f"  → Puntos en polilínea del eje:   {len(eje_coords)}")

    # ── c) Calcular longitud acumulada del eje ──
    acum = [0.0]
    for i in range(1, len(eje_coords)):
        d = haversine(*eje_coords[i-1], *eje_coords[i])
        acum.append(acum[-1] + d)
    longitud_total = acum[-1]
    print(f"  → Longitud total del eje: {longitud_total:.1f} m")

    # ── d) Estimar km de inicio del eje ──
    # El tramo empieza en KM 100+000 (100,000 m).
    # Buscamos el placemark más cercano a "100+000" para anclar.
    KM_INICIO = 99_960   # cadenamiento del primer punto del eje (metros)

    def km_a_coord(km_num):
        """Interpola coordenada en el eje dado un cadenamiento en metros."""
        dist_en_eje = km_num - KM_INICIO
        if dist_en_eje < 0:
            return eje_coords[0]
        if dist_en_eje >= longitud_total:
            return eje_coords[-1]
        # Búsqueda binaria
        lo, hi = 0, len(acum) - 1
        while lo < hi - 1:
            mid = (lo + hi) // 2
            if acum[mid] <= dist_en_eje:
                lo = mid
            else:
                hi = mid
        # Interpolar entre lo y hi
        seg_len = acum[hi] - acum[lo]
        if seg_len == 0:
            return eje_coords[lo]
        t = (dist_en_eje - acum[lo]) / seg_len
        lat = eje_coords[lo][0] + t * (eje_coords[hi][0] - eje_coords[lo][0])
        lon = eje_coords[lo][1] + t * (eje_coords[hi][1] - eje_coords[lo][1])
        return (lat, lon)

    # ── e) Asignar coordenadas a cada punto de control ──
    for pc in puntos_control:
        # Primero buscar por nombre exacto en placemarks
        coord = None
        for nombre, c in pm_coords.items():
            # Buscar coincidencia tipo "100+380" o "100380"
            km_str1 = f"{pc['km_num'] // 1000}+{pc['km_num'] % 1000:03d}"
            km_str2 = str(pc['km_num'])
            if km_str1 in nombre or km_str2 in nombre:
                coord = c
                break
        if coord is None:
            coord = km_a_coord(pc["km_num"])
        pc["lat"] = coord[0]
        pc["lon"] = coord[1]
        print(f"    {pc['label']:15s}  lat={pc['lat']:.7f}  lon={pc['lon']:.7f}")

    return puntos_control


# ─────────────────────────────────────────────
# 3. EXTRAER FOTOS Y COORDENADAS DEL TIMEMARK
# ─────────────────────────────────────────────

def extraer_fotos_timemark(kmz_path):
    """
    Devuelve lista de dicts:
      {name, lat, lon, img_bytes, img_name}
    """
    fotos = []
    with zipfile.ZipFile(kmz_path, "r") as z:
        # Leer KML
        kml_names = [n for n in z.namelist() if n.lower().endswith(".kml")]
        with z.open(kml_names[0]) as f:
            tree = ET.parse(f)
        root = tree.getroot()

        # Mapear archivos de imagen disponibles (insensible a mayúsculas)
        archivos_zip = {n.lower(): n for n in z.namelist()}

        for pm in root.iter("{http://www.opengis.net/kml/2.2}Placemark"):
            name_el  = pm.find("{http://www.opengis.net/kml/2.2}name")
            coord_el = pm.find(".//{http://www.opengis.net/kml/2.2}Point/"
                               "{http://www.opengis.net/kml/2.2}coordinates")
            if coord_el is None:
                continue

            raw = coord_el.text.strip().split(",")
            if len(raw) < 2:
                continue
            lon, lat = float(raw[0]), float(raw[1])

            nombre = name_el.text.strip() if name_el is not None else "sin_nombre"

            # Buscar imagen referenciada en <description> o <href>
            img_bytes = None
            img_name  = None

            desc_el = pm.find("{http://www.opengis.net/kml/2.2}description")
            href_candidates = []

            if desc_el is not None and desc_el.text:
                # Buscar src="..." o href="..." en el HTML de la descripción
                hrefs = re.findall(r'(?:src|href)=["\']([^"\']+)["\']',
                                   desc_el.text, re.IGNORECASE)
                href_candidates.extend(hrefs)

            # También buscar <IconStyle><Icon><href>
            for href_el in pm.iter("{http://www.opengis.net/kml/2.2}href"):
                if href_el.text:
                    href_candidates.append(href_el.text.strip())

            for href in href_candidates:
                # Normalizar ruta
                href_norm = href.replace("\\", "/").lstrip("/")
                href_low  = href_norm.lower()
                # Buscar en el zip
                real_name = archivos_zip.get(href_low)
                if real_name is None:
                    # Buscar solo por basename
                    base = os.path.basename(href_norm).lower()
                    real_name = next(
                        (v for k, v in archivos_zip.items()
                         if os.path.basename(k) == base), None
                    )
                if real_name:
                    img_bytes = z.read(real_name)
                    img_name  = real_name
                    break

            fotos.append({
                "name":      nombre,
                "lat":       lat,
                "lon":       lon,
                "img_bytes": img_bytes,
                "img_name":  img_name,
            })

    print(f"  → {len(fotos)} placemarks leídos del Timemark")
    con_foto = sum(1 for f in fotos if f["img_bytes"])
    print(f"  → {con_foto} con imagen adjunta")
    return fotos


# ─────────────────────────────────────────────
# 4. ASIGNAR FOTOS A PUNTOS DE CONTROL
# ─────────────────────────────────────────────

def asignar_fotos(fotos, puntos_control):
    """
    Asigna cada foto al punto de control más cercano (Haversine).
    Devuelve dict: label → [fotos]
    """
    grupos = {pc["label"]: [] for pc in puntos_control}

    for foto in fotos:
        mejor_label = None
        mejor_dist  = float("inf")
        for pc in puntos_control:
            d = haversine(foto["lat"], foto["lon"], pc["lat"], pc["lon"])
            if d < mejor_dist:
                mejor_dist  = d
                mejor_label = pc["label"]
        if mejor_label:
            grupos[mejor_label].append(foto)

    for pc in puntos_control:
        n = len(grupos[pc["label"]])
        if n:
            print(f"    {pc['label']:15s}  → {n} foto(s)")
    return grupos


# ─────────────────────────────────────────────
# 5. GENERAR EL WORD
# ─────────────────────────────────────────────

def generar_word(puntos_control, grupos, output_path):
    doc = Document()

    # ── Estilos generales ──
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Portada / título ──
    titulo = doc.add_heading("Registro Fotográfico por Punto de Control", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Tramo: Zacualtipan – Molango  |  KM 99+960 – 105+000\n"
        "Fuente: Timemark_202603222011.kmz"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    fotos_sin_imagen = 0

    for pc in puntos_control:
        label  = pc["label"]
        fotos  = grupos.get(label, [])
        if not fotos:
            continue

        # ── Encabezado de sección ──
        h = doc.add_heading(f"Punto de control {label}", level=2)
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        doc.add_paragraph(
            f"Coordenadas de referencia: "
            f"Lat {pc['lat']:.7f}  |  Lon {pc['lon']:.7f}\n"
            f"Total de fotografías: {len(fotos)}"
        )

        # ── Tabla 2 columnas ──
        cols = 2
        filas = math.ceil(len(fotos) / cols)
        tabla = doc.add_table(rows=filas * 2, cols=cols)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.style = "Table Grid"

        for idx, foto in enumerate(fotos):
            fila_img  = (idx // cols) * 2
            fila_cap  = fila_img + 1
            col_idx   = idx % cols

            celda_img = tabla.cell(fila_img, col_idx)
            celda_cap = tabla.cell(fila_cap, col_idx)

            # Imagen
            if foto["img_bytes"]:
                img_stream = io.BytesIO(foto["img_bytes"])
                try:
                    p = celda_img.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_stream, width=Inches(3.0))
                except Exception as e:
                    celda_img.text = f"[Error al insertar imagen: {e}]"
                    fotos_sin_imagen += 1
            else:
                celda_img.text = f"[Sin imagen: {foto['name']}]"
                fotos_sin_imagen += 1

            # Pie de foto
            cap_p = celda_cap.paragraphs[0]
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_p.add_run(
                f"{foto['name']}\n"
                f"Lat: {foto['lat']:.6f}  Lon: {foto['lon']:.6f}"
            )
            cap_run.font.size = Pt(8)
            cap_run.font.italic = True

        doc.add_paragraph()  # espacio entre secciones

    doc.save(output_path)
    print(f"\n✓ Word generado: {output_path}")
    if fotos_sin_imagen:
        print(f"  ⚠ {fotos_sin_imagen} foto(s) no pudieron insertarse (sin imagen en KMZ)")


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  Generador de Registro Fotográfico por Punto de Control")
    print("=" * 60)

    print("\n[1/4] Leyendo puntos de control del reporte geológico...")
    puntos = extraer_puntos_control_docx(DOCX_GEO)

    print("\n[2/4] Obteniendo coordenadas desde el KMZ de TRAZO...")
    puntos = coordenadas_desde_trazo(KMZ_TRAZO, puntos)

    print("\n[3/4] Extrayendo fotos del Timemark...")
    fotos = extraer_fotos_timemark(KMZ_TIMEMARK)

    print("\n[4/4] Asignando fotos a puntos de control...")
    grupos = asignar_fotos(fotos, puntos)

    print("\n[5/5] Generando documento Word...")
    generar_word(puntos, grupos, OUTPUT_DOCX)

    print("\nListo.")


if __name__ == "__main__":
    main()
